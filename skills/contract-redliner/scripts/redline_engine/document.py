"""LLM-facing view of a .docx contract.

A `DocumentView` is a frozen snapshot of a document for the LLM to reason
about: it carries paragraph IDs the model can cite, a Markdown serialization
that fits in a prompt, and enumerations of any pre-existing comments and
tracked changes (so the model can reply on threads and accept prior edits).

The view is built fresh each time the document is loaded; paragraph IDs are
assigned by load-order (`p-000`, `p-001`, ...) and are stable for the duration
of one redlining session. After we apply edits and save, reloading produces a
new view — the LLM must re-read in that case.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph
from lxml import etree

from .ooxml_utils import NSMAP, qn


_PARAGRAPH_ID_FMT = "p-{:03d}"


@dataclass
class ParagraphInfo:
    """A paragraph in the document, with the metadata the LLM and engine need."""

    id: str
    index: int
    text: str
    style: str
    is_empty: bool
    has_revisions: bool
    has_comments: bool
    # Word's auto-numbering is driven by w:numPr; we record the abstract
    # level (0 = top-level) when present so the serializer can indent.
    list_level: int | None = None


@dataclass
class ExistingComment:
    """A pre-existing comment in the document the LLM may need to reply to."""

    id: int
    author: str
    initials: str
    text: str
    paragraph_ids: list[str] = field(default_factory=list)
    parent_id: int | None = None  # if this comment is itself a reply


@dataclass
class ExistingRevision:
    """A pre-existing tracked change the LLM may need to accept/reject."""

    id: int
    type: str  # "ins" | "del" | "moveFrom" | "moveTo"
    author: str
    text: str
    paragraph_ids: list[str] = field(default_factory=list)


@dataclass
class DocumentView:
    docx_path: Path
    paragraphs: list[ParagraphInfo]
    comments: list[ExistingComment]
    revisions: list[ExistingRevision]
    # Map paragraph ID → underlying python-docx Paragraph object (not serialized
    # for the LLM; used internally when applying edits).
    _paragraph_objects: dict[str, Paragraph] = field(default_factory=dict, repr=False)
    _docx: DocxDocument | None = field(default=None, repr=False)

    @classmethod
    def load(cls, path: str | Path) -> "DocumentView":
        path = Path(path)
        doc = Document(str(path))
        paragraphs: list[ParagraphInfo] = []
        paragraph_objects: dict[str, Paragraph] = {}

        for i, p in enumerate(doc.paragraphs):
            pid = _PARAGRAPH_ID_FMT.format(i)
            info = ParagraphInfo(
                id=pid,
                index=i,
                text=p.text,
                style=p.style.name if p.style else "",
                is_empty=not p.text.strip(),
                has_revisions=_paragraph_has_revisions(p),
                has_comments=_paragraph_has_comments(p),
                list_level=_paragraph_list_level(p),
            )
            paragraphs.append(info)
            paragraph_objects[pid] = p

        comments = _enumerate_comments(doc, paragraphs, paragraph_objects)
        revisions = _enumerate_revisions(doc, paragraphs, paragraph_objects)

        return cls(
            docx_path=path,
            paragraphs=paragraphs,
            comments=comments,
            revisions=revisions,
            _paragraph_objects=paragraph_objects,
            _docx=doc,
        )

    def get(self, paragraph_id: str) -> ParagraphInfo | None:
        for p in self.paragraphs:
            if p.id == paragraph_id:
                return p
        return None

    def runtime_paragraph(self, paragraph_id: str) -> Paragraph | None:
        """Return the underlying python-docx `Paragraph` object for the given id."""
        return self._paragraph_objects.get(paragraph_id)

    def to_markdown(self, max_chars: int | None = None) -> str:
        """Serialize the document into an LLM-friendly Markdown view.

        Blank paragraphs are omitted from the visible output but their IDs are
        preserved (skipped) so paragraph IDs match the underlying document.

        If `max_chars` is set, truncates with a sentinel — useful for prompt
        budgeting. The LargeCo template is ~90 KB; most current frontier models
        comfortably fit the whole thing, but smaller models or budget runs may
        need windowing.
        """
        lines: list[str] = [f"# {self.docx_path.name}", ""]
        for p in self.paragraphs:
            if p.is_empty:
                continue
            flags = []
            if p.has_revisions:
                flags.append("REVS")
            if p.has_comments:
                flags.append("CMTS")
            flag_str = f" {{{','.join(flags)}}}" if flags else ""
            indent = "  " * (p.list_level or 0)
            lines.append(f"[{p.id}] ({p.style}){flag_str} {indent}{p.text}")

        if self.comments:
            lines.append("")
            lines.append("## Existing comments")
            for c in self.comments:
                head = f"[cmt-{c.id}] {c.author!r}"
                if c.parent_id is not None:
                    head += f" (reply to cmt-{c.parent_id})"
                spans = ",".join(c.paragraph_ids) or "?"
                lines.append(f"{head} on [{spans}]: {c.text}")

        if self.revisions:
            lines.append("")
            lines.append("## Existing tracked changes")
            for r in self.revisions:
                spans = ",".join(r.paragraph_ids) or "?"
                lines.append(
                    f"[rev-{r.id}] {r.type} by {r.author!r} on [{spans}]: {r.text!r}"
                )

        out = "\n".join(lines)
        if max_chars is not None and len(out) > max_chars:
            out = out[: max_chars - 100] + "\n\n[... document truncated ...]"
        return out

    def stats(self) -> dict[str, int]:
        return {
            "paragraphs": len(self.paragraphs),
            "non_empty_paragraphs": sum(1 for p in self.paragraphs if not p.is_empty),
            "existing_comments": len(self.comments),
            "existing_revisions": len(self.revisions),
        }


# --- Paragraph metadata helpers -------------------------------------------- #


def _paragraph_has_revisions(p: Paragraph) -> bool:
    for tag in ("w:ins", "w:del", "w:moveFrom", "w:moveTo"):
        if next(p._p.iter(qn(tag)), None) is not None:
            return True
    return False


def _paragraph_has_comments(p: Paragraph) -> bool:
    for tag in ("w:commentRangeStart", "w:commentRangeEnd", "w:commentReference"):
        if next(p._p.iter(qn(tag)), None) is not None:
            return True
    return False


def _paragraph_list_level(p: Paragraph) -> int | None:
    """Return the list nesting level if the paragraph is part of a numbered/bulleted list."""
    numPr = p._p.find(f"{qn('w:pPr')}/{qn('w:numPr')}")
    if numPr is None:
        return None
    ilvl = numPr.find(qn("w:ilvl"))
    if ilvl is None:
        return 0
    val = ilvl.get(qn("w:val"))
    try:
        return int(val) if val is not None else 0
    except ValueError:
        return 0


# --- Comment & revision enumeration ---------------------------------------- #


def _paragraph_ids_for_element_range(
    start_id_attr: str,
    end_id_attr: str,
    el_id: int,
    paragraphs: list[ParagraphInfo],
    paragraph_objects: dict[str, Paragraph],
) -> list[str]:
    """Find paragraph IDs that the given range marker pair spans.

    Walks paragraphs in document order looking for `w:commentRangeStart`/`End`
    (or analogous markers) with matching IDs. Returns paragraph IDs from start
    to end inclusive.
    """
    started = False
    spans: list[str] = []
    target = str(el_id)
    for pi in paragraphs:
        p = paragraph_objects[pi.id]
        # Did this paragraph contain start?
        for el in p._p.iter(qn(start_id_attr)):
            if el.get(qn("w:id")) == target:
                started = True
                break
        if started:
            spans.append(pi.id)
        # Did this paragraph contain end?
        for el in p._p.iter(qn(end_id_attr)):
            if el.get(qn("w:id")) == target:
                return spans
    return spans


def _enumerate_comments(
    doc: DocxDocument,
    paragraphs: list[ParagraphInfo],
    paragraph_objects: dict[str, Paragraph],
) -> list[ExistingComment]:
    out: list[ExistingComment] = []
    # python-docx exposes `Document.comments` (CommentCollection) in 1.2.
    try:
        collection = doc.comments
    except AttributeError:
        return out

    # Build a paraId → parent_id lookup from commentsExtended.xml so we can
    # mark replies. The commentsExtended part is optional; absence just means
    # no thread metadata available.
    paraid_parent = _read_comments_extended(doc)

    for c in collection:
        cid = int(c._element.get(qn("w:id"), -1))
        if cid < 0:
            continue
        spans = _paragraph_ids_for_element_range(
            "w:commentRangeStart",
            "w:commentRangeEnd",
            cid,
            paragraphs,
            paragraph_objects,
        )
        # Text — join all w:t descendants of the w:comment element.
        text = "".join(t.text or "" for t in c._element.iter(qn("w:t")))
        # Determine if this is a reply: look up our first paragraph's paraId.
        parent_id = None
        first_para = c._element.find(qn("w:p"))
        if first_para is not None:
            pid_attr = first_para.get(qn("w14:paraId"))
            if pid_attr and pid_attr in paraid_parent:
                parent_id = paraid_parent[pid_attr]

        out.append(
            ExistingComment(
                id=cid,
                author=c._element.get(qn("w:author"), ""),
                initials=c._element.get(qn("w:initials"), ""),
                text=text,
                paragraph_ids=spans,
                parent_id=parent_id,
            )
        )
    return out


def _read_comments_extended(doc: DocxDocument) -> dict[str, int]:
    """Return a paraId → parent_comment_id map from commentsExtended.xml.

    Word stores comment-thread parentage by paraId (a w14 attribute on the
    paragraph inside the comment body) rather than by comment IDs directly.
    We resolve that mapping here.
    """
    out: dict[str, int] = {}
    package_part = doc.part.package
    # Walk for any part whose URI ends with /word/commentsExtended.xml
    for rel in doc.part.rels.values():
        if rel.target_ref.endswith("commentsExtended.xml"):
            blob = rel.target_part.blob
            root = etree.fromstring(blob)
            # Build paraId → parent_paraId from <w15:commentEx>
            paraid_to_parent = {}
            for ex in root.iter(qn("w15:commentEx")):
                pid = ex.get(qn("w15:paraId"))
                ppid = ex.get(qn("w15:paraIdParent"))
                if pid:
                    paraid_to_parent[pid] = ppid
            if not paraid_to_parent:
                return out
            # Now we need to resolve parent paraId → comment id by walking
            # comments.xml.
            paraid_to_commentid = {}
            for crel in doc.part.rels.values():
                if crel.target_ref.endswith("comments.xml"):
                    cblob = crel.target_part.blob
                    croot = etree.fromstring(cblob)
                    for c in croot.iter(qn("w:comment")):
                        first_p = c.find(qn("w:p"))
                        if first_p is None:
                            continue
                        pid = first_p.get(qn("w14:paraId"))
                        cid = c.get(qn("w:id"))
                        if pid and cid is not None:
                            try:
                                paraid_to_commentid[pid] = int(cid)
                            except ValueError:
                                continue
            for pid, ppid in paraid_to_parent.items():
                if ppid in paraid_to_commentid:
                    out[pid] = paraid_to_commentid[ppid]
            break
    return out


def _enumerate_revisions(
    doc: DocxDocument,
    paragraphs: list[ParagraphInfo],
    paragraph_objects: dict[str, Paragraph],
) -> list[ExistingRevision]:
    out: list[ExistingRevision] = []
    seen_ids: set[int] = set()
    for pi in paragraphs:
        p = paragraph_objects[pi.id]
        for tag in ("w:ins", "w:del", "w:moveFrom", "w:moveTo"):
            for el in p._p.iter(qn(tag)):
                raw = el.get(qn("w:id"))
                if raw is None:
                    continue
                try:
                    rid = int(raw)
                except ValueError:
                    continue
                if rid in seen_ids:
                    continue
                seen_ids.add(rid)
                # Extract text — w:t (insertion) or w:delText (deletion)
                text_parts = []
                for t in el.iter(qn("w:t")):
                    text_parts.append(t.text or "")
                for t in el.iter(qn("w:delText")):
                    text_parts.append(t.text or "")
                out.append(
                    ExistingRevision(
                        id=rid,
                        type=etree.QName(el).localname,
                        author=el.get(qn("w:author"), ""),
                        text="".join(text_parts),
                        paragraph_ids=[pi.id],
                    )
                )
    return out
